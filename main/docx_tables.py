from docx.shared import Pt,Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

#Akt
def fix_table_layout(table):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr

    for el in tblPr.findall("w:tblLayout", tbl.nsmap):
        tblPr.remove(el)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

#Akt width
def set_col_width(cell, cm):
    twips = int(cm * 567)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for el in tcPr.findall("w:tcW", tcPr.nsmap):
        tcPr.remove(el)

    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

#Akt text
def set_cell_text(cell, text, bold=False, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("" if text is None else str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(7)
    run.bold = bold

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

#Akt
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    for el in tblPr.findall("w:tblBorders", tbl.nsmap):
        tblPr.remove(el)

    borders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)

    tblPr.append(borders)

#Akt
def force_tbl_grid(table, widths_cm):
    tbl = table._tbl

    # eski tblGrid ni o‘chiramiz
    for el in tbl.findall("w:tblGrid", tbl.nsmap):
        tbl.remove(el)

    tblGrid = OxmlElement("w:tblGrid")

    for w in widths_cm:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(w * 567)))
        tblGrid.append(gridCol)

    # ❗ tblGrid HAR DOIM tblPr DAN KEYIN turishi shart
    tbl.insert(1, tblGrid)


def replace_text(doc, replacements: dict):
    """
    DOCX ichidagi barcha paragraph va table cell run'larida
    matnlarni almashtiradi (Word + LibreOffice mos).
    """
    # Oddiy paragraphlar
    for p in doc.paragraphs:
        for run in p.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

    # Jadval ichidagi matnlar
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for old, new in replacements.items():
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                                run.font.name = "Times New Roman"
                                run.font.size = Pt(10)


def create_table_akt(doc, title, data, headers):

    widths = [1, 4, 3, 5, 2, 2, 4, 4, 3]

    # Title
    h = doc.add_paragraph()
    r = h.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    # ✅ 2 qatorli header
    table = doc.add_table(rows=2, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    fix_table_layout(table)
    set_table_borders(table)
    force_tbl_grid(table, widths)

    h1 = table.rows[0].cells
    h2 = table.rows[1].cells

    # 0..5 ustunlar (№..Birligi) - vertikal merge
    for col in range(0, 6):
        cell = h1[col].merge(h2[col])
        set_cell_text(cell, headers[col], bold=True, center=True)
        set_col_width(cell, widths[col])

    # ✅ Guruh: 6..8 (F.I.Sh, Lavozimi, Narxi)
    grp = h1[6].merge(h1[7]).merge(h1[8])
    set_cell_text(grp, "Kimga o'rnatildi yoki o'rnatilganligini tasdiqlovchi mas'ul shaxs (xarajatlar)", bold=True, center=True)

    # pastki headerlar
    set_cell_text(h2[6], headers[6], bold=True, center=True)  # F.I.Sh.
    set_cell_text(h2[7], headers[7], bold=True, center=True)  # Lavozimi
    set_cell_text(h2[8], headers[8], bold=True, center=True)  # Narxi

    set_col_width(h2[6], widths[6])
    set_col_width(h2[7], widths[7])
    set_col_width(h2[8], widths[8])

    # ✅ Data qatorlari
    for idx, row in enumerate(data, start=1):
        cells = table.add_row().cells

        full = [idx] + list(row)   # № qo‘shildi => 9 ta bo‘ladi
        while len(full) < 9:
            full.append("")

        for i, val in enumerate(full[:9]):
            set_cell_text(cells[i], val, center=True)
            set_col_width(cells[i], widths[i])

    return h, table


def create_table_akt_all(doc, title, data, headers):

    widths = [1, 4, 3, 4, 2, 2, 3, 2, 2, 2, 2]

    # Title
    h = doc.add_paragraph()
    r = h.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    # ✅ 2 qatorli header
    table = doc.add_table(rows=2, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    fix_table_layout(table)
    set_table_borders(table)
    force_tbl_grid(table, widths)

    h1 = table.rows[0].cells
    h2 = table.rows[1].cells

    # 0..5 ustunlar (№..Birligi) - vertikal merge
    for col in range(0, 6):
        cell = h1[col].merge(h2[col])
        set_cell_text(cell, headers[col], bold=True, center=True)
        set_col_width(cell, widths[col])

    # ✅ Guruh: 6..8 (F.I.Sh, Lavozimi, Narxi)
    grp = h1[6].merge(h1[7]).merge(h1[8]).merge(h1[9]).merge(h1[10])
    set_cell_text(grp, "Kimga o'rnatildi yoki o'rnatilganligini tasdiqlovchi mas'ul shaxs (xarajatlar)", bold=True, center=True)

    # pastki headerlar
    set_cell_text(h2[6], headers[6], bold=True, center=True)  # F.I.Sh.
    set_cell_text(h2[7], headers[7], bold=True, center=True)  # Lavozimi
    set_cell_text(h2[8], headers[8], bold=True, center=True)  # Narxi
    set_cell_text(h2[9], headers[9], bold=True, center=True)  # Id
    set_cell_text(h2[10], headers[10], bold=True, center=True)  # Sana

    set_col_width(h2[6], widths[6])
    set_col_width(h2[7], widths[7])
    set_col_width(h2[8], widths[8])
    set_col_width(h2[9], widths[9])
    set_col_width(h2[10], widths[10])

    # ✅ Data qatorlari
    for idx, row in enumerate(data, start=1):
        cells = table.add_row().cells

        full = [idx] + list(row)   # № qo‘shildi => 9 ta bo‘ladi
        while len(full) < 11:
            full.append("")

        for i, val in enumerate(full[:11]):
            set_cell_text(cells[i], val, center=True)
            set_col_width(cells[i], widths[i])

    return h, table


def set_column_widths(table, widths_cm):
    """
    Ustun kengliklarini sm bo‘yicha o‘rnatish.
    widths_cm: [1, 3.5, 3, 2] kabi ro‘yxat (sm).
    """
    for col, w in zip(table.columns, widths_cm):
        col.width = Cm(w)


def style_cell_paragraph(cell, bold=False, center=True, font_size=11):
    """
    Har bir katak matnini normal ko‘rinishga keltirish:
    - Times New Roman
    - 11 pt
    - bo‘sh joylarsiz
    """
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(font_size)
            r.bold = bold

#Dallatnoma table
def create_table(doc, title, data, headers):
    if not data:
        return None, None

    # 🔥 WORD-FIRST WIDTHS (sm)
    if len(headers) == 4:
        widths = [1, 6, 4, 4]
    else:
        widths = [1, 7, 4]

    # --- Sarlavha ---
    h = doc.add_paragraph()
    r = h.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 🔥 MUHIM SOZLAMALAR
    fix_table_layout(table)
    set_table_borders(table)
    force_tbl_grid(table, widths)

    # --- Header ---
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        style_cell_paragraph(hdr[i], bold=True, center=True)
        set_col_width(hdr[i], widths[i])

    # --- Data ---
    for idx, row in enumerate(data, start=1):
        tr = table.add_row().cells
        tr[0].text = str(idx)
        tr[1].text = row.get("name") or ""

        if len(headers) == 4:
            tr[2].text = row.get("serial") or ""
            tr[3].text = row.get("inventory") or ""
        else:
            tr[2].text = row.get("serial") or ""

        for i, cell in enumerate(tr):
            style_cell_paragraph(cell, center=True)
            set_col_width(cell, widths[i])

    return h, table

#Svod table
def create_table_cols_svod(doc, data, headers, grand_total=0):
    widths = [1, 7.5, 2, 2, 4, 4, 5.5, 2]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    fix_table_layout(table)
    set_table_borders(table)
    force_tbl_grid(table, widths)

    # Header
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_text(hdr[i], text, bold=True, center=True)
        set_col_width(hdr[i], widths[i])

    # Data
    for idx, row in enumerate(data, start=1):
        cells = table.add_row().cells
        full = [idx] + list(row)

        while len(full) < len(headers):
            full.append("")

        for i, val in enumerate(full[:len(headers)]):
            set_cell_text(cells[i], val, center=True)
            set_col_width(cells[i], widths[i])

    # ✅ 1 ta JAMI qator (0..5 merge, summa 5-ustunda)
    sum_value = f"{grand_total:,}".replace(",", " ")

    r = table.add_row().cells

    # 0..4 merge qilsak, 5-ustun (Umumiy qiymati) alohida qoladi
    m = r[0]
    for j in range(1, 5):   # 0..4 merge
        m = m.merge(r[j])

    set_cell_text(r[0], "J A M I:", bold=True, center=True)
    set_cell_text(r[5], sum_value, bold=True, center=True)

    # qolgan ustunlar bo'sh qoladi
    set_cell_text(r[6], "", center=False)  # Eslatma
    set_cell_text(r[7], "", center=True)   # Kod 1C

    return table

#Reestr text
def set_cell_text_reestr(cell, text, bold=False, center=False, font_size=6.5):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("" if text is None else str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.bold = bold

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


#Reestr table
def create_table_cols_reestr(doc, data, grand_total=0):
    """
    data: har bir qator 14 ta qiymat (№ bu funksiyada qo‘shiladi)
    [Qurilma, Seriya, Material, Soni, Birlik narx, Umumiy,
     FIO, Lavozim, Tashkilot, Kim o‘rnatgan, Sana, Sorov №, Sorov sana, 1C]
    """

    # 15 ta ustun uchun width ham 15 ta bo‘lsin
    widths = [0.7, 2.2, 2, 3.0, 0.9, 1.7, 1.9, 2.7, 2.0, 3.0, 2.7, 1.7, 1.4, 1.7, 1.4]

    # ✅ 2 qatorli header, 15 ustun
    table = doc.add_table(rows=2, cols=15)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    fix_table_layout(table)
    set_table_borders(table)
    force_tbl_grid(table, widths)

    h1 = table.rows[0].cells
    h2 = table.rows[1].cells

    # Bitta ustunli sarlavhalar (yuqoridan pastga merge)
    single = {
        0: "№",
        1: "Qurilma nomi",
        2: "Seriya №",
        3: "Sarf materiallari nomi",
        4: "Soni",
        5: "Birlikdagi narxi",
        6: "Materiallarning\numumiy qiymati",
        9: "Tashkilot, bo'lim nomi",
        10: "Kim tomonidan\no'rnatilgan",
        11: "O'rnatish\nsanasi",
        12: "So'rovnoma №",
        13: "So'rovnoma\nsanasi",
        14: "1C\nkodi",
    }

    for col_idx, text in single.items():
        cell = h1[col_idx].merge(h2[col_idx])
        set_cell_text_reestr(cell, text, bold=True, center=True, font_size=6.5)
        set_col_width(cell, widths[col_idx])

    # Guruhlangan header: Qurilmadan foydalanuvchi (7,8)
    grp = h1[7].merge(h1[8])
    set_cell_text_reestr(grp, "Qurilmadan foydalanuvchi", bold=True, center=True, font_size=6.5)

    set_cell_text_reestr(h2[7], "F.I.Sh.", bold=True, center=True, font_size=6.5)
    set_cell_text_reestr(h2[8], "Lavozimi", bold=True, center=True, font_size=6.5)

    set_col_width(h2[7], widths[7])
    set_col_width(h2[8], widths[8])

    # ✅ Data qatorlari
    for idx, row in enumerate(data, start=1):
        cells = table.add_row().cells

        # row: 14 ta bo‘lishi kerak, № boshiga qo‘shiladi => 15
        full = [idx] + list(row)

        while len(full) < 15:
            full.append("")

        for i, val in enumerate(full[:15]):
            set_cell_text_reestr(cells[i], val, center=True, font_size=6.5)
            set_col_width(cells[i], widths[i])

    # ✅ JAMI qatori: summa 6-ustunda (Umumiy qiymat)
    sum_value = f"{int(grand_total or 0):,}".replace(",", " ")
    r = table.add_row().cells

    # 0..5 merge (№ dan Birlik narxigacha)
    merged = r[0]
    for j in range(1, 6):   # 1..5
        merged = merged.merge(r[j])

    set_cell_text_reestr(r[0], "J A M I:", bold=True, center=True, font_size=7)
    set_cell_text_reestr(r[6], sum_value, bold=True, center=True, font_size=7)

    # qolgan ustunlar bo‘sh
    for k in range(7, 15):
        set_cell_text_reestr(r[k], "", center=True, font_size=7)

    return table


import os
import subprocess
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.colors import red
from PyPDF2 import PdfReader, PdfWriter
import qrcode
import shutil
import base64
import json
from django.conf import settings
from django.utils import timezone
from contextlib import contextmanager
from qrcode.constants import ERROR_CORRECT_M

# ==========================================================
# 1) DOCX → PDF (LIBREOFFICE)  → (pdf_path, debug)
# ==========================================================
def convert_docx_to_pdf_libre(docx_path: str) -> tuple[str | None, str]:
    if not os.path.exists(docx_path):
        return None, f"DOCX topilmadi: {docx_path}"

    output_dir = os.path.dirname(docx_path)
    expected_pdf = os.path.splitext(docx_path)[0] + ".pdf"

    # ✅ soffice aniqlash (tez + to'g'ri)
    if os.name == "nt":
        candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        soffice = next((c for c in candidates if os.path.exists(c)), None)
    else:
        soffice = shutil.which("soffice") or shutil.which("libreoffice")

    if not soffice:
        return None, "LibreOffice (soffice) topilmadi. (sudo apt install libreoffice)"

    env = os.environ.copy()
    env.setdefault("HOME", output_dir)

    cmd = [
        soffice,
        "--headless",
        "--invisible",
        "--nodefault",
        "--nologo",
        "--nofirststartwizard",
        "--norestore",
        "--convert-to", "pdf",
        "--outdir", output_dir,
        docx_path
    ]

    try:
        proc = subprocess.run(
            cmd,
            stdout=subprocess.DEVNULL,   # ✅ stdout/stderr ni o'qimaslik tezroq
            stderr=subprocess.DEVNULL,
            env=env,
            timeout=30                  # ✅ osilib qolmasin
        )
    except Exception as e:
        return None, str(e)

    if proc.returncode != 0:
        return None, f"LibreOffice returncode={proc.returncode}"

    if os.path.exists(expected_pdf):
        return expected_pdf, "OK"

    # LibreOffice ba'zan nomni o'zgartirishi mumkin → fallback
    pdfs = [f for f in os.listdir(output_dir) if f.lower().endswith(".pdf")]
    if pdfs:
        return os.path.join(output_dir, pdfs[0]), "OK (fallback)"

    return None, "PDF yaratilmadi"



@contextmanager
def file_lock(lock_path: str, timeout: int = 10):
    start = timezone.now()
    while True:
        try:
            fd = os.open(lock_path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
            os.close(fd)
            break
        except FileExistsError:
            if (timezone.now() - start).total_seconds() > timeout:
                raise TimeoutError("PDF band (lock timeout)")
    try:
        yield
    finally:
        try:
            os.remove(lock_path)
        except Exception:
            pass


def create_overlay_pdf(page_w: float, page_h: float, text: str, qr_link: str, overlay_path: str):
    qr_png = overlay_path.replace(".pdf", "_qr.png")

    # 🔥 QR ni qo‘lda yaratamiz
    qr = qrcode.QRCode(
        version=None,
        error_correction=ERROR_CORRECT_M,
        box_size=6,
        border=1,
    )
    qr.add_data(qr_link)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")
    img.save(qr_png)

    c = canvas.Canvas(overlay_path, pagesize=(page_w, page_h))
    c.setFillColor(red)

    # Yozuv (yuqori chap)
    c.setFont("Helvetica-Bold", 10)
    c.drawString(10, page_h - 15, text)

    # QR (pastki markaz)
    qr_size = 60
    x_center = (page_w - qr_size) / 2
    c.drawImage(qr_png, x_center, 2, width=qr_size, height=qr_size, mask="auto")

    c.save()

    try:
        os.remove(qr_png)
    except Exception:
        pass


from django.urls import reverse
def sign_pdf_inplace(pdf_path: str, request, approver_name: str, deed_id: int) -> None:
    """
    PDF ning o'zini o'ziga imzolaydi: HAR BIR SAHIFAGA QR + yozuv qo'yadi.
    QR skaner bo‘lsa: /deed/status/<deed_id>/ ga olib boradi.
    """

    if not pdf_path or not os.path.exists(pdf_path):
        raise FileNotFoundError("PDF topilmadi")

    if not pdf_path.lower().endswith(".pdf"):
        raise ValueError("PDF noto‘g‘ri")

    abs_pdf = os.path.abspath(pdf_path)

    # ✅ QR endi STATUS sahifaga olib boradi
    qr_link = request.build_absolute_uri(
        reverse("deed_status", args=[int(deed_id)])
    )

    # Matn
    text = (
        f"Ushbu hujjat {approver_name} tomonidan "
        f"{timezone.now().strftime('%Y-%m-%d %H:%M')} da tasdiqlandi."
    )

    lock_path = abs_pdf + ".lock"
    with file_lock(lock_path, timeout=10):
        reader = PdfReader(abs_pdf)
        if not reader.pages:
            raise ValueError("PDF bo‘sh")

        first = reader.pages[0]
        w = float(first.mediabox.width)
        h = float(first.mediabox.height)

        overlay_path = abs_pdf.replace(".pdf", "_overlay.pdf")
        tmp_out = abs_pdf.replace(".pdf", "_signed_tmp.pdf")

        create_overlay_pdf(w, h, text, qr_link, overlay_path)

        overlay_reader = PdfReader(overlay_path)
        overlay_page = overlay_reader.pages[0]

        writer = PdfWriter()
        for page in reader.pages:
            page.merge_page(overlay_page)  # ✅ har bir sahifaga
            writer.add_page(page)

        with open(tmp_out, "wb") as f:
            writer.write(f)

        shutil.move(tmp_out, abs_pdf)

        try:
            os.remove(overlay_path)
        except Exception:
            pass


# ==========================================================
# 3) Original + Overlay PDF birlashtirish
# ==========================================================
def merge_pdf(original: str, overlay: str, output: str) -> None:
    reader = PdfReader(original)
    overlay_reader = PdfReader(overlay)
    writer = PdfWriter()

    overlay_page = overlay_reader.pages[0]

    for page in reader.pages:
        page.merge_page(overlay_page)
        writer.add_page(page)

    with open(output, "wb") as f:
        writer.write(f)


# =========================================================
# 4) Asosiy FUNKSIYA: DOCX → PDF → Overlay → Signed PDF
# ==========================================================
import shutil

def sign_pdf(pdf_path: str, request, approver_name: str) -> bool:

    if not os.path.exists(pdf_path):
        print("❌ PDF topilmadi:", pdf_path)
        return False

    pdf_path = os.path.abspath(pdf_path)

    text = (
        f"Ushbu hujjat {approver_name} tomonidan "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')} da tasdiqlandi."
    )

    media_root = os.path.abspath(settings.MEDIA_ROOT)
    rel_pdf = os.path.relpath(pdf_path, media_root).replace(os.sep, "/")

    # QR doim shu faylni ochadi (fayl nomi o‘zgarmaydi)
    qr_link = request.build_absolute_uri(settings.MEDIA_URL + rel_pdf)

    overlay_path = pdf_path.replace(".pdf", "_overlay.pdf")
    merged_tmp = pdf_path.replace(".pdf", "_merged_tmp.pdf")

    create_overlay_pdf(
        original_pdf_path=pdf_path,
        text=text,
        qr_link=qr_link,
        overlay_path=overlay_path
    )

    merge_pdf(
        original=pdf_path,
        overlay=overlay_path,
        output=merged_tmp
    )

    # 🔥 ASOSIY NUQTA
    # Imzolangan hujjat — eski faylning o‘ziga yoziladi
    shutil.move(merged_tmp, pdf_path)

    if os.path.exists(overlay_path):
        os.remove(overlay_path)

    return True



def decode_jwt(token):
    payload = token.split(".")[1]
    payload += "=" * (-len(payload) % 4)
    return json.loads(base64.urlsafe_b64decode(payload))


def get_sso_redirect_uri(request):
    host = request.get_host()
    if "localhost" in host or "127.0.0.1" in host:
        return "http://localhost:8000/sso/callback/"
    return "https://report.imv.uz/sso/callback/"

def short_fio(emp):
    """Zayniddinov S. F. ko'rinishiga keltiradi."""
    if not emp:
        return ""
    ln = (getattr(emp, "last_name", "") or "").strip()
    fn = (getattr(emp, "first_name", "") or "").strip()
    mn = (getattr(emp, "father_name", "") or "").strip()
    fi = (fn[:1] + ".") if fn else ""
    mi = (mn[:1] + ".") if mn else ""
    return f"{ln} {fi} {mi}".replace("  ", " ").strip()


def insert_receivers_into_placeholder(doc, placeholder: str, receivers):
    """
    reestr.docx ichidagi 'RECEIVER' placeholder turgan joyga
    receiverlar ro'yxatini qator-qator qilib yozadi.
    """
    p = next((p for p in doc.paragraphs if placeholder in p.text), None)
    if not p:
        return False

    # placeholder matnni tozalash
    p.text = ""
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    receivers_list = list(receivers)
    if not receivers_list:
        p.add_run("—")
        return True

    for idx, emp in enumerate(receivers_list, start=1):
        # chap tomonda rank yoki "1-toifali mutaxassis" chiqishini xohlasangiz:
        left = emp.rank.name if getattr(emp, "rank", None) else f"{idx}-toifali mutaxassis"
        right = short_fio(emp)

        line = f"{left}  {right}"

        if idx == 1:
            p.add_run(line)
        else:
            np = p.insert_paragraph_after(line)
            np.paragraph_format.space_before = Pt(0)
            np.paragraph_format.space_after = Pt(0)
            np.paragraph_format.line_spacing = 1
            np.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p = np

    return True


def short_fio(emp):
    """Zayniddinov S. F. ko'rinishiga keltiradi."""
    if not emp:
        return ""
    ln = (getattr(emp, "last_name", "") or "").strip()
    fn = (getattr(emp, "first_name", "") or "").strip()
    mn = (getattr(emp, "father_name", "") or "").strip()
    fi = (fn[:1] + ".") if fn else ""
    mi = (mn[:1] + ".") if mn else ""
    return f"{ln} {fi} {mi}".replace("  ", " ").strip()


def insert_receivers_into_placeholder(doc, placeholder: str, receivers):
    """
    reestr.docx ichidagi 'RECEIVER' placeholder turgan joyga
    receiverlar ro'yxatini qator-qator qilib yozadi.
    """
    p = next((p for p in doc.paragraphs if placeholder in p.text), None)
    if not p:
        return False

    # placeholder matnni tozalash
    p.text = ""
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    receivers_list = list(receivers)
    if not receivers_list:
        p.add_run("—")
        return True

    for idx, emp in enumerate(receivers_list, start=1):
        # chap tomonda rank yoki "1-toifali mutaxassis" chiqishini xohlasangiz:
        left = emp.rank.name if getattr(emp, "rank", None) else f"{idx}-toifali mutaxassis"
        right = short_fio(emp)

        line = f"{left}  {right}"

        if idx == 1:
            p.add_run(line)
        else:
            np = p.insert_paragraph_after(line)
            np.paragraph_format.space_before = Pt(0)
            np.paragraph_format.space_after = Pt(0)
            np.paragraph_format.line_spacing = 1
            np.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p = np

    return True
